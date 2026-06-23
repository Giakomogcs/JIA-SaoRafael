#!/usr/bin/env python3
"""
Conversor Markdown -> Word (.docx) com identidade visual SENAI.
Processa DOCUMENTACAO.md e gera Documentacao_SaoRafael.docx.

Recursos:
  - Capa estilizada (faixa vermelha SENAI)
  - Cabecalho/rodape com numero de pagina
  - Sumario (TOC) atualizavel
  - Titulos coloridos (H1 vermelho, H2 com regua, H3/H4)
  - Tabelas zebradas com cabecalho vermelho
  - Callouts (NOTE/TIP/WARNING/DANGER) com barra lateral
  - Blocos de codigo com fundo escuro
  - Mermaid -> PNG (se mmdc disponivel) ou bloco de codigo com legenda
  - Markdown inline (**negrito**, *italico*, `codigo`)
"""

import os
import re
import shutil
import hashlib
import subprocess
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paleta SENAI ───────────────────────────────────────────────
ACCENT   = RGBColor(0xE3, 0x06, 0x13)   # vermelho SENAI
ACCENT2  = RGBColor(0x8B, 0x04, 0x10)   # vermelho escuro
INK      = RGBColor(0x0B, 0x0B, 0x0B)   # titulos escuros
INK2     = RGBColor(0x1F, 0x29, 0x37)   # corpo
GRAY     = RGBColor(0x6B, 0x72, 0x80)   # legenda
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
CODE_FG  = RGBColor(0xE5, 0xE7, 0xEB)

HEX_ACCENT  = 'E30613'
HEX_RULE    = 'D1D5DB'
HEX_ZEBRA   = 'F9F7F7'
HEX_CODEBG  = '0B0B0B'
HEX_INLINE  = 'EEF2FF'

CALLOUTS = {
    'NOTE':    ('1D4ED8', 'EFF4FF', 'i', 'Nota'),
    'INFO':    ('1D4ED8', 'EFF4FF', 'i', 'Informacao'),
    'TIP':     ('047857', 'ECFDF5', '+', 'Dica'),
    'WARNING': ('B45309', 'FFFBEB', '!', 'Atencao'),
    'DANGER':  ('B91C1C', 'FEF2F2', 'x', 'Critico'),
    'CAUTION': ('B91C1C', 'FEF2F2', 'x', 'Cuidado'),
}

FONT = 'Calibri'
MONO = 'Consolas'


# ─── helpers de baixo nivel (XML) ───────────────────────────────
def _set_shading(element, hex_fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    element.append(shd)


def shade_cell(cell, hex_fill):
    _set_shading(cell._tc.get_or_add_tcPr(), hex_fill)


def set_cell_borders(cell, hex_color=HEX_RULE, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), hex_color)
        borders.append(el)
    tcPr.append(borders)


def set_left_bar(paragraph, hex_color, sz=24):
    """Barra lateral grossa (callout)."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(sz))
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), hex_color)
    pbdr.append(left)
    pPr.append(pbdr)


def set_bottom_rule(paragraph, hex_color=HEX_ACCENT, sz=8):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), hex_color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_field(paragraph, instr, color=None, size=None):
    """Insere campo do Word (ex.: PAGE, NUMPAGES, TOC)."""
    r = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr_el = OxmlElement('w:instrText'); instr_el.set(qn('xml:space'), 'preserve'); instr_el.text = instr
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_begin); r._r.append(instr_el); r._r.append(fld_sep); r._r.append(fld_end)
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    r.font.name = FONT
    return r


# ─── markdown inline ────────────────────────────────────────────
INLINE_RE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))')


def add_inline(paragraph, text, base_color=INK2, base_size=11):
    """Renderiza **negrito**, *italico*, `codigo`, [link](url)."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            _run(paragraph, text[pos:m.start()], base_color, base_size)
        tok = m.group(0)
        if tok.startswith('**'):
            _run(paragraph, tok[2:-2], base_color, base_size, bold=True)
        elif tok.startswith('`'):
            r = _run(paragraph, tok[1:-1], INK, base_size - 1, mono=True)
            shade_run(r, HEX_INLINE)
        elif tok.startswith('['):
            label = re.match(r'\[([^\]]+)\]', tok).group(1)
            _run(paragraph, label, ACCENT2, base_size, bold=True)
        elif tok.startswith('*'):
            _run(paragraph, tok[1:-1], base_color, base_size, italic=True)
        pos = m.end()
    if pos < len(text):
        _run(paragraph, text[pos:], base_color, base_size)


def _run(paragraph, text, color, size, bold=False, italic=False, mono=False):
    r = paragraph.add_run(text)
    r.font.name = MONO if mono else FONT
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return r


def shade_run(run, hex_fill):
    rPr = run._r.get_or_add_rPr()
    _set_shading(rPr, hex_fill)


# ─── blocos de conteudo ─────────────────────────────────────────
def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        _run(p, text.upper(), ACCENT, 20, bold=True)
        set_bottom_rule(p, HEX_ACCENT, 12)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        _run(p, text, INK, 15, bold=True)
        set_bottom_rule(p, HEX_RULE, 6)
    elif level == 3:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, ACCENT2, 13, bold=True)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        _run(p, text, INK2, 11.5, bold=True)
    return p


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, HEX_CODEBG)
    first = True
    for ln in code.split('\n'):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(ln if ln else ' ')
        r.font.name = MONO
        r.font.size = Pt(9)
        r.font.color.rgb = CODE_FG
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc, kind, lines):
    color, fill, icon, label = CALLOUTS[kind]
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    head = cell.paragraphs[0]
    set_left_bar(head, color, 28)
    head.paragraph_format.space_after = Pt(2)
    _run(head, f'{icon}  {label}', RGBColor.from_string(color), 11, bold=True)
    for ln in lines:
        p = cell.add_paragraph()
        set_left_bar(p, color, 28)
        p.paragraph_format.space_after = Pt(1)
        add_inline(p, ln, INK2, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, rows):
    headers = [c.strip() for c in rows[0].strip().strip('|').split('|')]
    data = []
    for row in rows[2:]:
        if row.strip():
            cells = [c.strip() for c in row.strip().strip('|').split('|')]
            data.append(cells)
    if not headers:
        return
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # cabecalho
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        shade_cell(cell, HEX_ACCENT)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        add_inline(p, h, WHITE, 10.5)
        for rn in p.runs:
            rn.bold = True
    # dados
    for i, drow in enumerate(data):
        zebra = (i % 2 == 1)
        for j in range(len(headers)):
            cell = table.rows[i + 1].cells[j]
            set_cell_borders(cell)
            if zebra:
                shade_cell(cell, HEX_ZEBRA)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            val = drow[j] if j < len(drow) else ''
            add_inline(p, val, INK2, 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ─── mermaid ────────────────────────────────────────────────────
def find_mmdc():
    for name in ('mmdc.cmd', 'mmdc'):
        path = shutil.which(name)
        if path:
            return path
    return None


MMDC = find_mmdc()
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CACHE = os.path.join(BASE_DIR, '.cache')
PUPPETEER_CFG = os.path.join(BASE_DIR, 'puppeteer-config.json')


def render_mermaid(code):
    if not MMDC:
        return None
    os.makedirs(CACHE, exist_ok=True)
    h = hashlib.sha1(code.encode('utf-8')).hexdigest()[:12]
    out = os.path.join(CACHE, f'{h}.png')
    if os.path.exists(out):
        return out
    src = os.path.join(CACHE, f'{h}.mmd')
    with open(src, 'w', encoding='utf-8') as f:
        f.write(code)
    cmd = [MMDC, '-i', src, '-o', out, '-b', 'white', '-w', '1600', '-s', '2']
    if os.path.exists(PUPPETEER_CFG):
        cmd += ['-p', PUPPETEER_CFG]
    try:
        subprocess.run(cmd, check=True, capture_output=True, shell=(os.name == 'nt'))
        return out if os.path.exists(out) else None
    except Exception:
        return None


def add_mermaid(doc, code, idx):
    png = render_mermaid(code)
    if png:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(png, width=Inches(6.2))
        except Exception:
            add_code_block(doc, code)
    else:
        add_code_block(doc, code)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(cap, f'Diagrama {idx} (Mermaid)', GRAY, 9, italic=True)


# ─── capa, header, footer ───────────────────────────────────────
def build_cover(doc, title, author):
    band = doc.add_table(rows=1, cols=1)
    bcell = band.cell(0, 0)
    shade_cell(bcell, HEX_ACCENT)
    bp = bcell.paragraphs[0]
    bp.paragraph_format.space_before = Pt(10)
    bp.paragraph_format.space_after = Pt(10)
    _run(bp, 'SENAI', WHITE, 24, bold=True)
    _run(bp, '  ·  Documentacao Tecnica', WHITE, 14)

    for _ in range(4):
        doc.add_paragraph()

    pt = doc.add_paragraph()
    _run(pt, title, INK, 30, bold=True)

    ps = doc.add_paragraph()
    _run(ps, 'Documento de Negocio e Tecnico', ACCENT2, 14, bold=True)

    for _ in range(8):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    set_left_bar(meta, HEX_ACCENT, 24)
    _run(meta, author, GRAY, 11)
    meta2 = doc.add_paragraph()
    set_left_bar(meta2, HEX_ACCENT, 24)
    _run(meta2, f'Gerado em {date.today().strftime("%d/%m/%Y")}', GRAY, 11)

    doc.add_page_break()


def build_toc(doc):
    h = doc.add_paragraph()
    _run(h, 'SUMARIO', ACCENT, 18, bold=True)
    set_bottom_rule(h, HEX_ACCENT, 12)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', INK2, 11)
    note = doc.add_paragraph()
    _run(note, '(Clique com o botao direito > Atualizar campo para gerar o indice.)', GRAY, 8.5, italic=True)
    doc.add_page_break()


def setup_header_footer(doc, title, project):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ''
    tab = OxmlElement('w:tabs'); t = OxmlElement('w:tab')
    t.set(qn('w:val'), 'right'); t.set(qn('w:pos'), '9360'); tab.append(t)
    hp._p.get_or_add_pPr().append(tab)
    _run(hp, title, GRAY, 8.5)
    hp.add_run('\t')
    _run(hp, project, ACCENT2, 8.5, bold=True)
    set_bottom_rule(hp, HEX_RULE, 4)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ''
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fp, 'Pagina ', GRAY, 9)
    add_field(fp, 'PAGE', GRAY, 9)
    _run(fp, ' de ', GRAY, 9)
    add_field(fp, 'NUMPAGES', GRAY, 9)


# ─── parser principal ───────────────────────────────────────────
def parse(md_file, docx_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    title = 'Documentacao Tecnica'
    author = 'Doc Master'
    project = 'Sao Rafael'

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK2

    lines = content.split('\n')
    i = 0

    # frontmatter
    if lines and lines[0].strip() == '---':
        i = 1
        while i < len(lines) and lines[i].strip() != '---':
            m = re.match(r'(\w+):\s*"?(.+?)"?\s*$', lines[i])
            if m:
                if m.group(1) == 'title':
                    title = m.group(2)
                elif m.group(1) == 'author':
                    author = m.group(2)
            i += 1
        i += 1

    setup_header_footer(doc, title, project)
    build_cover(doc, title, author)
    build_toc(doc)

    mermaid_idx = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # code / mermaid fence
        if stripped.startswith('```'):
            lang = stripped[3:].strip().lower()
            buf = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                buf.append(lines[i]); i += 1
            i += 1
            if lang == 'mermaid':
                mermaid_idx += 1
                add_mermaid(doc, '\n'.join(buf), mermaid_idx)
            else:
                add_code_block(doc, '\n'.join(buf))
            continue

        # callout: > [!TYPE]
        cm = re.match(r'>\s*\[!(\w+)\]', stripped)
        if cm and cm.group(1).upper() in CALLOUTS:
            kind = cm.group(1).upper()
            buf = []
            i += 1
            while i < len(lines) and lines[i].lstrip().startswith('>'):
                buf.append(re.sub(r'^\s*>\s?', '', lines[i])); i += 1
            add_callout(doc, kind, [b for b in buf if b.strip()])
            continue

        # blockquote generico -> callout NOTE
        if stripped.startswith('>'):
            buf = []
            while i < len(lines) and lines[i].lstrip().startswith('>'):
                buf.append(re.sub(r'^\s*>\s?', '', lines[i])); i += 1
            add_callout(doc, 'NOTE', [b for b in buf if b.strip()])
            continue

        # table
        if '|' in line and stripped.startswith('|'):
            buf = [line]
            i += 1
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                buf.append(lines[i]); i += 1
            if len(buf) >= 2:
                add_table(doc, buf)
            continue

        # headings
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            add_heading(doc, stripped.lstrip('#').strip(), min(level, 4))
            i += 1
            continue

        # horizontal rule
        if stripped in ('---', '***', '___'):
            sep = doc.add_paragraph()
            set_bottom_rule(sep, HEX_RULE, 4)
            i += 1
            continue

        # bullets
        if re.match(r'^[-*]\s+', stripped):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(1)
            add_inline(p, stripped[2:], INK2, 10.5)
            i += 1
            continue

        # numbered
        if re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(1)
            add_inline(p, re.sub(r'^\d+\.\s+', '', stripped), INK2, 10.5)
            i += 1
            continue

        # empty
        if not stripped:
            i += 1
            continue

        # paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_inline(p, stripped, INK2, 11)
        i += 1

    doc.save(docx_file)
    print(f'OK -> {docx_file}')
    print(f'Mermaid renderizado: {"sim" if MMDC else "nao (mmdc ausente; usado bloco de codigo)"}')


if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    md = os.path.join(base, 'DOCUMENTACAO.md')
    out = os.path.join(base, 'Documentacao_SaoRafael.docx')
    print('Convertendo com identidade visual SENAI...')
    parse(md, out)
    print('Concluido.')
